from fileinput import filename
import time
import streamlit as st
import json
import openai, pandas as pd
import fitz  # PyMuPDF
from pydantic import BaseModel, Field, ValidationError
from typing import List, Dict, Union
import base64
from docx import Document  
from docx.shared import Pt
from io import BytesIO
import re, ast

# Initialize AzureChatOpenAI
azure_endpoint = "https://chat-gpt-a1.openai.azure.com/"
api_key = "c09f91126e51468d88f57cb83a63ee36"
deployment_name = "DanielChatGPT16k"

openai.api_type = "azure"
openai.api_key = api_key
openai.api_base = azure_endpoint
openai.api_version = "2024-02-15-preview"

class TrademarkDetails(BaseModel):
    trademark_name: str = Field(description="The name of the Trademark", example="DISCOVER")
    status: str = Field(description="The Status of the Trademark", example="Registered")
    serial_number: str = Field(description="The Serial Number of the trademark from Chronology section", example="87−693,628")
    international_class_number: List[int] = Field(description="The International class number or Nice Classes number of the trademark from Goods/Services section or Nice Classes section", example=[18])
    owner: str = Field(description="The owner of the trademark", example="WALMART STORES INC")
    goods_services: str = Field(description="The goods/services from the document", example="LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES")
    page_number: int = Field(description="The page number where the trademark details are found in the document", example=3)
    registration_number: Union[str, None] = Field(description="The Registration number of the trademark from Chronology section", example="5,809,957")
    design_phrase: str = Field(description="The design phrase of the trademark", example="THE MARK CONSISTS OF THE STYLIZED WORD 'MINI' FOLLOWED BY 'BY MOTHERHOOD.'", default="")
    

def preprocess_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'[\u2013\u2014]', '-', text)
    return text

def is_correct_format_code1(page_text: str) -> bool:
    required_fields = ["Status:", "Goods/Services:"] # , "Last Reported Owner:"
    return all(field in page_text for field in required_fields)

def is_correct_format_code2(page_text: str) -> bool:
    required_fields = ["Register", "Nice Classes", "Goods & Services"]
    return all(field in page_text for field in required_fields)

def extract_trademark_details_code1(document_chunk: str) -> Dict[str, Union[str, List[int]]]:
    response = openai.ChatCompletion.create(
        engine=deployment_name,
        messages=[
            {"role": "system", "content": "You are a helpful assistant for extracting Meta Data from the Trademark Document."},
            {"role": "user", "content": f"Extract the following details from the trademark document: trademark name, status, serial number, international class number, owner, goods & services, filed date, registration number.\n\nDocument:\n{document_chunk}"}
        ],
        max_tokens=4000,
        temperature=0
    )

    extracted_text = response.choices[0].message['content'].strip()
    details = {}
    for line in extracted_text.split("\n"):
        if ":" in line:
            key, value = line.split(":", 1)
            details[key.strip().lower().replace(" ", "_")] = value.strip()
    return details

def extract_registration_number(document: str) -> str:
    """ Extract the registration number from the Chronology section """
    match = re.search(r'Chronology:.*?Registration Number:\s*([\d,]+)', document, re.DOTALL)
    if match:
        return match.group(1).strip()
    return "No registration number presented in document"

def extract_trademark_details_code2(page_text: str) -> Dict[str, Union[str, List[int]]]:
    details = {}

    trademark_name_match = re.search(r"\d+\s*/\s*\d+\s*\n\s*\n\s*([A-Za-z0-9'&!,\-. ]+)\s*\n", page_text)
    if trademark_name_match:
        details["trademark_name"] = trademark_name_match.group(1).strip()
    else:
        trademark_name_match = re.search(r"(?<=\n)([A-Za-z0-9'&!,\-. ]+)(?=\n)", page_text)
        details["trademark_name"] = trademark_name_match.group(1).strip() if trademark_name_match else ""

    status_match = re.search(r'Status\s*(?:\n|:\s*)([A-Za-z]+)', page_text, re.IGNORECASE)
    details["status"] = status_match.group(1).strip() if status_match else ""

    owner_match = re.search(r'Holder\s*(?:\n|:\s*)(.*)', page_text, re.IGNORECASE)
    if owner_match:
        details["owner"] = owner_match.group(1).strip()
    else:
        owner_match = re.search(r'Owner\s*(?:\n|:\s*)(.*)', page_text, re.IGNORECASE)
        details["owner"] = owner_match.group(1).strip() if owner_match else ""
        
        

    nice_classes_match = re.search(r'Nice Classes\s*[\s:]*\n((?:\d+(?:,\s*\d+)*)\b)', page_text, re.IGNORECASE)
    if nice_classes_match:
        nice_classes_text = nice_classes_match.group(1)
        nice_classes = [int(cls.strip()) for cls in nice_classes_text.split(",")]
        details["international_class_number"] = nice_classes
    else:
        details["international_class_number"] = []
        


    serial_number_match = re.search(r'Application#\s*(.*)', page_text, re.IGNORECASE)
    details["serial_number"] = serial_number_match.group(1).strip() if serial_number_match else ""

    goods_services_match = re.search(r'Goods & Services\s*(.*?)(?=\s*G&S translation|$)', page_text, re.IGNORECASE | re.DOTALL)
    details["goods_services"] = goods_services_match.group(1).strip() if goods_services_match else ""
    
    registration_number_match = re.search(r'Registration#\s*(.*)', page_text, re.IGNORECASE)
    details["registration_number"] = registration_number_match.group(1).strip() if registration_number_match else ""
    
    # Description
    design_phrase = re.search(r'Description\s*(.*?)(?=\s*Applicant|Owner|Holder|$)', page_text, re.IGNORECASE | re.DOTALL)
    details["design_phrase"] = design_phrase.group(1).strip() if design_phrase else "No Design phrase presented in document"
    

    return details

def read_pdf(file_path: str, exclude_header_footer: bool = True) -> str:
    document_text = ""
    with fitz.open(file_path) as pdf_document:
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            if exclude_header_footer:
                rect = page.rect
                x0 = rect.x0
                y0 = rect.y0 + rect.height * 0.1
                x1 = rect.x1
                y1 = rect.y1 - rect.height * 0.1
                page_text = page.get_text("text", clip=(x0, y0, x1, y1))
            else:
                page_text = page.get_text()
            document_text += page_text
    return document_text

def split_text(text: str, max_tokens: int = 1500) -> List[str]:
    chunks = []
    current_chunk = []
    current_length = 0

    for line in text.split('\n'):
        line_length = len(line.split())
        if current_length + line_length > max_tokens:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [line]
            current_length = line_length
        else:
            current_chunk.append(line)
            current_length += line_length

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks

def parse_international_class_numbers(class_numbers: str) -> List[int]:
    numbers = class_numbers.split(',')
    return [int(num.strip()) for num in numbers if num.strip().isdigit()]

def extract_international_class_numbers_and_goods_services(document: str) -> Dict[str, Union[List[int], str]]:
    """ Extract the International Class Numbers and Goods/Services from the document """
    class_numbers = []
    goods_services = []
    pattern = r'International Class (\d+): (.*?)(?=\nInternational Class \d+:|\n[A-Z][a-z]+:|\nLast Reported Owner:|\Z)'
    matches = re.findall(pattern, document, re.DOTALL)
    for match in matches:
        class_number = int(match[0])
        class_numbers.append(class_number)
        goods_services.append(f"Class {class_number}: {match[1].strip()}")
    return {
        "international_class_numbers": class_numbers,
        "goods_services": "\n".join(goods_services)
    }
    
def extract_design_phrase(document: str) -> str:
    """ Extract the design phrase from the document """
    match = re.search(r'Design Phrase:\s*(.*?)(?=Other U\.S\. Registrations:|$)', document, re.DOTALL)
    if match:
        design_phrase = match.group(1).strip()
        # Remove any newline characters within the design phrase
        design_phrase = ' '.join(design_phrase.split())
        return design_phrase
    return "No Design phrase presented in document"

    
def parse_trademark_details(document_path: str) -> List[Dict[str, Union[str, List[int]]]]:
    with fitz.open(document_path) as pdf_document:
        all_extracted_data = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()
            
            if is_correct_format_code1(page_text):
                preprocessed_chunk = preprocess_text(page_text)
                extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                additional_data = extract_international_class_numbers_and_goods_services(page_text)
                registration_number = extract_registration_number(page_text)
        
                design_phrase = ""
                next_page_num = page_num + 1
                if next_page_num < pdf_document.page_count:
                    next_page = pdf_document.load_page(next_page_num)
                    next_page_text = next_page.get_text()
                    design_phrase = extract_design_phrase(next_page_text)
                
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    extracted_data.update(additional_data)
                    extracted_data["design_phrase"] = design_phrase
                    all_extracted_data.append(extracted_data)
                    extracted_data["registration_number"] = registration_number
                    
                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_name = data.get("trademark_name", "").split(',')[0].strip()
                        if "Global Filings" in trademark_name:
                            trademark_name = trademark_name.split("Global Filings")[0].strip()
                        owner = data.get("owner", "").split(',')[0].strip()
                        status = data.get("status", "").split(',')[0].strip()
                        serial_number = data.get("serial_number", "")
                        international_class_number = data.get("international_class_numbers", [])
                        goods_services = data.get("goods_services", "")
                        page_number = data.get("page_number", "")
                        registration_number = data.get("registration_number", "No registration number presented in document")
                        design_phrase = data.get("design_phrase", "No Design phrase presented in document")

                        # If crucial fields are missing, attempt to re-extract the values
                        if not trademark_name or not owner or not status or not international_class_number:
                            preprocessed_chunk = preprocess_text(data.get("raw_text", ""))
                            extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                            trademark_name = extracted_data.get("trademark_name", trademark_name).split(',')[0].strip()
                            if "Global Filings" in trademark_name:
                                trademark_name = trademark_name.split("Global Filings")[0].strip()
                            owner = extracted_data.get("owner", owner).split(',')[0].strip()
                            status = extracted_data.get("status", status).split(',')[0].strip()
                            international_class_number = parse_international_class_numbers(extracted_data.get("international_class_number", "")) or international_class_number
                            registration_number = extracted_data.get("registration_number", registration_number).split(',')[0].strip()

                        trademark_details = TrademarkDetails(
                            trademark_name=trademark_name,
                            owner=owner,
                            status=status,
                            serial_number=serial_number,
                            international_class_number=international_class_number,
                            goods_services=goods_services,
                            page_number=page_number,
                            registration_number=registration_number,
                            design_phrase=design_phrase
                        )
                        trademark_info = {
                            "trademark_name": trademark_details.trademark_name,
                            "owner": trademark_details.owner,
                            "status": trademark_details.status,
                            "serial_number": trademark_details.serial_number,
                            "international_class_number": trademark_details.international_class_number,
                            "goods_services": trademark_details.goods_services,
                            "page_number": trademark_details.page_number,
                            "registration_number":trademark_details.registration_number,
                            "design_phrase": trademark_details.design_phrase
                        }
                        trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")
                                    
            else :
                if not is_correct_format_code2(page_text):
                    continue

                extracted_data = extract_trademark_details_code2(page_text)
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_details = TrademarkDetails(
                            trademark_name=data.get("trademark_name", ""),
                            owner=data.get("owner", ""),
                            status=data.get("status", ""),
                            serial_number=data.get("serial_number", ""),
                            international_class_number=data.get("international_class_number", []),
                            goods_services=data.get("goods_services", ""),
                            page_number=data.get("page_number", 0),
                            registration_number=data.get("registration_number", ""),
                            design_phrase=data.get("design_phrase", "")
                        )
                        if (trademark_details.trademark_name != "" and trademark_details.owner != "" and trademark_details.status != "" and trademark_details.goods_services != ""):
                                trademark_info = {
                                    "trademark_name": trademark_details.trademark_name,
                                    "owner": trademark_details.owner,
                                    "status": trademark_details.status,
                                    "serial_number": trademark_details.serial_number,
                                    "international_class_number": trademark_details.international_class_number,
                                    "goods_services": trademark_details.goods_services,
                                    "page_number": trademark_details.page_number,
                                    "registration_number":trademark_details.registration_number,
                                    "design_phrase":trademark_details.design_phrase,
                                }
                                trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")

        return trademark_list

def compare_trademarks(existing_trademark: List[Dict[str, Union[str, List[int]]]], proposed_name: str, proposed_class: str, proposed_goods_services: str) -> List[Dict[str, int]]:
    proposed_classes = [int(c.strip()) for c in proposed_class.split(',')]
    response_reasoning = openai.ChatCompletion.create(
        engine=deployment_name,
        messages=[
            {"role": "system", "content": """You are a trademark attorney to properly reasoning based on given conditions and assign conflict grade high or moderate or low to existing trademark and respond with only High or Moderate or Low. \n\n 
                                            Conditions for determining Conflict Grades:\n\n 
                                            
                                            Condition 1: Trademark Name Comparison\n 
                                            - Condition 1A: The existing trademark's name is a character-for-character match with the proposed trademark name.\n 
                                            - Condition 1B: The existing trademark's name is semantically equivalent to the proposed trademark name.\n 
                                            - Condition 1C: The existing trademark's name is phonetically equivalent to the proposed trademark name.\n 
                                            - Condition 1D: Primary Position Requirement- In the context of trademark conflicts, the primary position of a trademark refers to the first word or phrase element in a multi-word or phrase trademark. For a conflict to exist between an existing trademark and a proposed trademark based on Condition 1D, the proposed trademark name must be in the primary position of the existing trademark. This means that the proposed trademark name should be the first word of the existing trademark.\n
                                                            Example:\n Existing Trademark: "STORIES AND JOURNEYS"\n Proposed Trademark: "JOURNEY"\n Analysis:\n The existing trademark "STORIES AND JOURNEYS" consists of multiple words/phrases.\n For the proposed trademark "JOURNEY" to be in conflict under Condition 1D, it must appear as the first word/phrase in the existing trademark.\n In this case, the first word/phrase in "STORIES AND JOURNEYS" is "STORIES", not "JOURNEY".\n Therefore, "JOURNEY" does not meet Condition 1D because it is not in the primary position of the existing trademark.\n
                                                            Example:\n Existing Trademark: "JOURNEY BY COMPANION"\n Proposed Trademark: "JOURNEY"\n Analysis:\n The existing trademark "JOURNEY BY COMPANION" consists of multiple words/phrases.\n For the proposed trademark "JOURNEY" to be in conflict under Condition 1D, it must appear as the first word/phrase in the existing trademark.\n In this case, the first word/phrase in "JOURNEY BY COMPANION" is "JOURNEY".\n Therefore, "JOURNEY" does meet Condition 1D because it is in the primary position of the existing trademark.\n
                                                                                                        
                                            Condition 2: Goods/Services Classification\n 
                                            - Condition 2: The existing trademark's goods/services are in the same class as those of the proposed trademark.\n
                                            
                                            Condition 3: Target Market and Products\n 
                                            - Condition 3A: The existing trademark's goods/services target the exact same products as the proposed trademark.\n 
                                            - Condition 3B: The existing trademark's goods/services target an exact market as the proposed trademark.\n
                                            
                                            If existing trademark in user given input satisfies:\n\n
                                            - Special case: If existing Trademark Status is Cancelled or Abandoned, they will automatically be considered as conflict grade low but still give the reasoning for the potential conflicts.\n\n
                                            - If the existing trademark satisfies Condition 1A, 1B, or 1C, and also satisfies the revised Condition 1D (when applicable), along with Condition 2, and both Condition 3A and 3B, then the conflict grade should be High.\n
                                            - If the existing trademark satisfies any two of the following: Condition 1A, 1B, or 1C (with the revised Condition 1D being a necessary component for these to be considered satisfied when applicable), Condition 2, Condition 3A and 3B, then the conflict grade should be Moderate.\n
                                            - If the existing trademark satisfies only one (or none) of the conditions: Condition 1A, 1B, 1C (only if the revised Condition 1D is also satisfied when applicable), Condition 2, Condition 3A and 3B, then the conflict grade should be Low.\n\n
                                            
                                            Format of the Response:\n
                                            Resoning for Conflict: Reasoning for conflict in bullet points (In reasoning, if exact same goods, services and industries: list the overlaps, you should reasoning whether the good/services are overlapping or not including classes (if same as proposed trademark or not) and see trademark name whether identical (character-for-character) matches, phonetic equivalents, if it is in primary position (first word in the phrase) or not, if it is not in primary position (first word in the phrase) of the existing trademark it is not conflicting and standard plural forms for subject goods and goods that may be related or not. Reasoning should be based on provided information. Do not provide any kind of hypothetical reasoning.)\n\n
                                            
                                            Step 1: Identifying Potential Conflicts
                                            - What is the existing trademark?
                                            - What is the status of the existing trademark?
                                            - Who is the owner of the existing trademark?
                                            - What is the class number for the existing trademark?
                                            - What is the proposed trademark?
                                            - Who is the applicant for the proposed trademark?
                                            - What is the class number for the proposed trademark?

                                            Step 2: Condition 1A - Character-for-Character Match
                                            - Does the existing trademark match the proposed trademark character-for-character?
                                            - If not, does the existing trademark form part of the proposed trademark? Specify the matching part.

                                            Step 3: Condition 1B - Sound-Alike Match
                                            - Do the existing and proposed trademarks sound alike when spoken?
                                            - If yes, describe the similarities in their pronunciation.

                                            Step 4: Condition 1C - Visual Similarity
                                            - Do the existing and proposed trademarks look visually similar?
                                            - If yes, describe the visual elements that contribute to this similarity.

                                            Step 5: Condition 1D - Primary Position
                                            - Is the existing trademark in the primary position of the proposed trademark?
                                            - Is the primary position the beginning of the proposed trademark?
                                            - Does the proposed trademark exactly match the existing trademark in its primary position?
                                            - If the proposed trademark name is a single word or phrase, it must be fully presented in the primary position of the existing trademark for this condition to be applicable.
                                            - As if the existing trademark name is a single word, then condition 1D is not applicable.

                                            Step 6: Condition 2 - Class Overlap
                                            - Do the class numbers of the existing and proposed trademarks match?
                                            - If the class numbers are different, are the classes related in a way that might cause consumer confusion?

                                            Step 7: Condition 3A - Goods/Services Overlap
                                            - What goods/services and products are covered by the existing trademark?
                                            - What goods/services and products are covered by the proposed trademark?
                                            - Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?

                                            Step 8: Condition 3B - Target Market
                                            - Who is the target market for the goods/services covered by the existing trademark?
                                            - Who is the target market for the goods/services covered by the proposed trademark?
                                            - Is there an exact match or exact overlap in the target market for the existing and proposed trademarks?

                                            Step 9: Conflict Grade Assessment
                                            - Based on the analysis, how would you grade the potential conflict?
                                            - Provide a brief reasoning for the conflict grade.
                                            
                                            Example Analysis Using the Steps :
                                            - Trademark Name: MH
                                            - Trademark Status: REGISTERED
                                            - Trademark Owner: ZHAO
                                            - Trademark Class Number: 3
                                            - Proposed Trademark: MH BY MOTHERHOOD
                                            - Applicant: ABC Company
                                            - Proposed Trademark Class Number: 3

                                            Step 2: Condition 1A - Character-for-Character Match
                                            - Does the existing trademark match the proposed trademark character-for-character?
                                            - No, "MH" is part of "MH BY MOTHERHOOD" but not an exact match.

                                            Step 3: Condition 1B - Sound-Alike Match
                                            - Do the existing and proposed trademarks sound alike when spoken?
                                            - Yes, The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are phonetically equivalent
                                            - If the existing trademark name is multi-word, then the primary word of the existing trademark name followed by a color or the name of a country, city, or number (either numeric or letter) should be considered as Condition satisfied.

                                            Step 4: Condition 1C - Visual Similarity
                                            - Do the existing and proposed trademarks look visually similar?
                                            - No, "MH" and "MH BY MOTHERHOOD" do not look visually similar.

                                            Step 5: Condition 1D - Primary Position
                                            - Is the existing trademark in the primary position of the proposed trademark?
                                            - Yes, the proposed trademark "MH BY MOTHERHOOD", which has "MH" as in the primary position of the existing trademark "MH".

                                            Step 6: Condition 2 - Class Overlap
                                            - Do the class numbers of the existing and proposed trademarks match?
                                            - Yes, both are in Class 3.

                                            Step 7: Condition 3A - Goods/Services Overlap
                                            - What goods/services are covered by the existing trademark?
                                            - Non-medicated cosmetic and hair care preparations.
                                            - What goods/services are covered by the proposed trademark?
                                            - Hair care preparations.
                                            - Is there an overlap between the goods/services and product of the existing and proposed trademarks?
                                            - Yes, both include exact products such as hair care preparations.

                                            Step 8: Condition 3B - Target Market
                                            - Who is the target market for the goods/services covered by the existing trademark?
                                            - Consumers interested in cosmetic and hair care products.
                                            - Who is the target market for the goods/services covered by the proposed trademark?
                                            - Consumers interested in hair care products.
                                            - Is there an overlap in the target market for the existing and proposed trademarks?
                                            - Yes, both target the exact same consumers interested in hair care products.

                                            Step 9: Conflict Grade Assessment
                                            - Based on the analysis, how would you grade the potential conflict?
                                            - High.

                                            Provide a brief reasoning for the conflict grade.
                                            - Example : Condition 1A is not satisfied since there is no character-for-character match and 1C are not satisfied either. However, Conditions 1B and Condition 1D is been satisfied and also there is a class overlap and goods/services overlap, as well as a shared target market, indicating a High potential for consumer confusion.
                                                                                    
                                            Conflict Grade: Based on above reasoning (Low or Moderate or High)."""
                                            },
            
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: DISCOVER WHAT’S NEXT\n
                                            Goods/Services: Online retail store services geared for men and women, featuring a wide variety of unique consumer products\n 
                                            International Class Numbers: 35\n
                                            Status: REGISTERED\n
                                            Owner: THE GIDDYUP GROUP, INC\n
                                            
                                            Proposed Trademark:\n
                                            Name: DISCOVER\n 
                                            Goods/Services: Luggage and carrying bags; suitcases, trunks, travelling bags, sling bags for = carrying infants, school bags; purses; wallets; retail and online retail services\n
                                            International Class Numbers: 18, 35\n"""
            },
        {"role": "assistant", "content":"""
Reasoning for Conflict:
Step 1: Condition 1A - Exact Character Match
- Does the existing trademark contain the proposed trademark term?
- Yes, the existing trademark "DISCOVER WHAT’S NEXT" contains the term "DISCOVER," which is a character for-character match with the proposed trademark "DISCOVER."
- Condition 1A is satisfied.

Step 2: Condition 1D - Primary Position
- Is the existing trademark in the primary position of the proposed trademark?
- Yes, "DISCOVER" is in the primary position of the existing trademark "DISCOVER WHAT’S NEXT."
- As if the existing trademark name is a single word, then condition 1D is not applicable.
- Existing trademark name is a single word, Hence Condition 1D is not applicable.

Step 3: Condition 2 - Class Overlap
- Do the class numbers of the existing and proposed trademarks overlap?
- Yes, the existing trademark is registered under International Class 35, which overlaps with the proposed trademark's Class 35 for retail and online retail services.
- Condition 2 is satisfied.

Step 4: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include online retail store services geared for men and women, featuring a wide variety of unique consumer products.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include luggage and carrying bags; suitcases, trunks, travelling bags, sling bags for carrying infants, school bags; purses; wallets; retail and online retail services.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- No, while both trademarks include retail and online retail services, the existing trademark focuses on a wide variety of unique consumer products, whereas the proposed trademark focuses specifically on luggage, bags, and related items.
- Condition 3A is not fully satisfied.

Step 5: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in a wide variety of unique consumer products.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in luggage, carrying bags, and related products.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- No, although there is some overlap in the general market of retail and online retail services, the specific focus of the products differs.
- Condition 3B is not fully satisfied.

Conclusion:
- Despite satisfying Conditions 1A, (1D not applicable) , and 2, the differences in the specific goods/services and target markets mean that the conflict is not high.
- Since Conditions 3A and 3B are not fully satisfied due to the lack of exact overlap in the goods/services and target markets, the conflict grade should be moderate rather than high.

Reason:
Reasoning for Conflict:
1A - The existing trademark "DISCOVER WHAT’S NEXT" contains the term "DISCOVER," which is a character for-character match with the proposed trademark "DISCOVER."
1D - "DISCOVER" is in the primary position of the existing trademark "DISCOVER WHAT’S NEXT". However, - Existing trademark name is a single word, Hence Condition 1D is not applicable. 
2  - The existing trademark is registered under International Class 35, which overlaps with the proposed trademark's Class 35 for retail and online retail services.
3A - Both trademarks include retail and online retail services, the existing trademark focuses on a wide variety of unique consumer products, whereas the proposed trademark focuses specifically on luggage, bags, and related items.
3B - Although there is some overlap in the general market of retail and online retail services, the specific focus of the products differs.

Conclusion:
- Despite satisfying Conditions 1A, (1D not applicable) , and 2, the differences in the specific goods/services and target markets mean that the conflict is not high.
- Since Conditions 3A and 3B are not fully satisfied due to the lack of exact overlap in the goods/services and target markets, the conflict grade should be moderate rather than high.

- Conflict Grade: Moderate

"""
            },
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: DB JOURNEY\n
                                            Goods/Services: All-purpose carrying bags, rucksacks, hipsacks, luggage, toiletry bags, key bags, luggage tags, pocket wallets, straps for luggage, shoulder straps, and umbrellas (Class 18)\n 
                                            International Class Numbers: 9, 16, 18, 25, 28\n
                                            Status: PENDING SECTION 66(A) (MADRID PROTOCOL)\n
                                            Owner: DB EQUIPMENT AS NORWAY AS\n
                                            
                                            Proposed Trademark:\n
                                            Name: JOURNEY\n 
                                            Goods/Services: Luggage and carrying bags; suitcases, trunks, travelling bags, sling bags for carrying infants, school bags; purses; wallets; retail and online retail services\n
                                            International Class Numbers: 18, 35\n"""
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "DB JOURNEY" is not a character-for-character match with the proposed trademark "JOURNEY."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are phonetically similar due to the shared term "JOURNEY."
- Condition 1C is satisfied.

Step 4: Condition 1D - Primary Position
- The term "JOURNEY" is in the primary position in the proposed trademark "JOURNEY."
- The term "JOURNEY" is not in the primary position in the existing trademark "DB JOURNEY" (the primary term is "DB").
- Condition 1D is not satisfied.

Step 5: Condition 2 - Class Overlap
- The existing trademark includes Class 18, which overlaps with the proposed trademark's Class 18 for luggage and carrying bags.
- The existing trademark does not include Class 35, which is part of the proposed trademark's classification.
- Condition 2 is partially satisfied.

Step 6: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services in Class 18 include all-purpose carrying bags, rucksacks, hipsacks, luggage, toiletry bags, key bags, luggage tags, pocket wallets, straps for luggage, shoulder straps, and umbrellas.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services in Class 18 include luggage and carrying bags, suitcases, trunks, travelling bags, sling bags for carrying infants, school bags, purses, and wallets.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- There is a significant overlap in the goods/services in Class 18.
- Condition 3A is satisfied.

Step 7: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in a wide range of bags and carrying cases, including luggage and related accessories.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in luggage, carrying bags, and related products.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- There is an overlap in the target market for the existing and proposed trademarks.
- Condition 3B is satisfied.

Conclusion:
- Conditions 1C, 2 (partially), 3A, and 3B are satisfied.
- Condition 1D is not satisfied due to the proposed trademark name "JOURNEY" not being in the primary position of the existing trademark name "DB JOURNEY."
- Given the significant overlap in goods/services and the shared target market but the non-satisfaction of - Condition 1D, the conflict grade should be moderate rather than high.

Reason:
Reasoning for Conflict:
1A - The existing trademark "DB JOURNEY" is not a character-for-character match with the proposed trademark "JOURNEY."
1B - The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are not semantically equivalent.
1C - The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are phonetically similar due to the shared term "JOURNEY."
1D - The term "JOURNEY" is not in the primary position in the existing trademark "DB JOURNEY" (the primary term is "DB").
2  - The existing trademark includes Class 18, which overlaps with the proposed trademark's Class 18 for luggage and carrying bags.
3A - There is a significant overlap in the goods/services in Class 18.
3B - There is an overlap in the target market for the existing and proposed trademarks.

Conclusion:
- Conditions 1C, 2 (partially), 3A, and 3B are satisfied.
- Condition 1D is not satisfied due to the proposed trademark name "JOURNEY" not being in the primary position of the existing trademark name "DB JOURNEY."
- Given the significant overlap in goods/services and the shared target market but the non-satisfaction of - Condition 1D, the conflict grade should be moderate rather than high.

- Conflict Grade: Moderate

"""
            },
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: RADIANT RED\n
                                            Goods/Services: Non−medicated hair care preparations\n 
                                            International Class Numbers: 3\n
                                            Status: REGISTERED\n
                                            Owner: KAO KABUSHIKI KAISHA TA KAO CORPORATION JAPAN CORPORATION\n
                                            
                                            Proposed Trademark:\n
                                            Name: RADIANT AMBER\n 
                                            Goods/Services: DEODORANT\n
                                            International Class Numbers: 3\n"""
            },
            {"role": "assistant", "content":"""  
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "RADIANT RED" is not a character-for-character match with the proposed trademark "RADIANT AMBER."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "RADIANT RED" and the proposed trademark "RADIANT AMBER" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "RADIANT RED" and the proposed trademark "RADIANT AMBER" are phonetically similar due to the shared term "RADIANT."
- Condition 1C is satisfied.

Step 4: Condition 1D - Primary Position
- The term "RADIANT" is in the primary position in the proposed trademark "RADIANT AMBER."
- The term "RADIANT" is in the primary position in the existing trademark "RADIANT RED."
- Both the existing trademark and the proposed trademark have "RADIANT" in their primary positions.
- If the proposed trademark name is a single word or phrase, it must be fully presented in the primary position of the existing trademark for this condition to be applicable.
- If the existing trademark name is a single word, then condition 1D is not applicable.
- Condition 1D is satisfied.

Step 5: Condition 2 - Class Overlap
- Both the existing and proposed trademarks are in International Class 3.
- Condition 2 is satisfied.

Step 6: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include non-medicated hair care preparations.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include deodorant.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- No, there is no exact match or overlap as the existing trademark covers hair care preparations and the proposed trademark covers deodorant.
- Condition 3A is not satisfied.

Step 7: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in non-medicated hair care preparations.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in deodorant.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- No, there is no exact overlap in the target market as they focus on different types of personal care products.
- Condition 3B is not satisfied.

Conclusion:
- Conditions 1C, 1D, and 2 are satisfied.
- Conditions 1A, 1B, 3A, and 3B are not satisfied.
- Given the phonetic similarity, primary position match, and class overlap but the significant differences in goods/services and target markets, the conflict grade should be moderate.

Reason:
1A - The existing trademark "RADIANT RED" is not a character-for-character match with the proposed trademark "RADIANT AMBER."
1B - The existing trademark "RADIANT RED" and the proposed trademark "RADIANT AMBER" are not semantically equivalent.
1C - The existing trademark "RADIANT RED" and the proposed trademark "RADIANT AMBER" are phonetically similar due to the shared term "RADIANT."
1D - Both the existing trademark and the proposed trademark have "RADIANT" in their primary positions.
2  - Both trademarks are in International Class 3.
3A - There is no exact match or overlap between the goods/services in Class 3.
3B - There is no exact overlap in the target market for the existing and proposed trademarks.

Conclusion:
- Conditions 1C, 1D, and 2 are satisfied.
- Conditions 1A, 1B, 3A, and 3B are not satisfied.
- Given the phonetic similarity, primary position match, and class overlap but the significant differences in goods/services and target markets, the conflict grade should be moderate.

- Conflict Grade: Moderate             

            """},
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: MH\n
                                            Goods/Services: Non−medicated cosmetic hair care preparations in the nature of hair wax; cosmetic hair filling fibers for covering bald and thinning spots on the scalp\n 
                                            International Class Numbers: 3\n
                                            Status: REGISTERED\n
                                            Owner: ZHAO\n
                                            
                                            Proposed Trademark:\n
                                            Name: MH BY MOTHERHOOD\n 
                                            Goods/Services: IC 003: SKIN CARE PREPARATIONS; COSMETICS; BABY CARE PRODUCTS, NAMELY, SKIN SOAPS, BABY WASH, BABY BUBBLE BATH, BABY LOTIONS, BABY SHAMPOOS; SKIN CLEANSERS; BABY WIPES; NON− MEDICATED DIAPER RASH OINTMENTS AND LOTIONS; SKIN LOTIONS, CREAMS, MOISTURIZERS, AND OILS; BODY WASH; BODY SOAP; DEODORANTS; PERFUME; HAIR CARE PREPARATIONS\n
                                            International Class Numbers: 3\n"""
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "MH" is not a character-for-character match with the proposed trademark "MH BY MOTHERHOOD."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are phonetically similar due to the shared term "MH."
- Condition 1C is satisfied.

Step 4: Condition 1D - Primary Position
- The term "MH" is in the primary position in the proposed trademark "MH BY MOTHERHOOD."
- The term "MH" is in the primary position in the existing trademark "MH."
- As if the existing trademark name is a single word, then condition 1D is not applicable.
- Existing trademark name is a single word, Hence Condition 1D is not applicable.
- Condition 1D is not applicable.

Step 5: Condition 2 - Class Overlap
- Both the existing and proposed trademarks are in International Class 3.
- Condition 2 is satisfied.

Step 6: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include non-medicated cosmetic hair care preparations in the nature of hair wax and cosmetic hair filling fibers for covering bald and thinning spots on the scalp.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include skin care preparations, cosmetics, baby care products, skin cleansers, baby wipes, non-medicated diaper rash ointments and lotions, skin lotions, creams, moisturizers, oils, body wash, body soap, deodorants, perfume, and hair care preparations.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- There is a partial overlap in the goods/services, specifically in the area of hair care preparations.
- Condition 3A is partially satisfied.

Step 7: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in non-medicated cosmetic hair care preparations.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in a broader range of skin care, cosmetics, baby care products, and hair care preparations.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- There is a partial overlap in the target market, specifically for consumers interested in hair care preparations.
- Condition 3B is partially satisfied.

Conclusion:
- Conditions 1C, 2, and (1D not applicable) are satisfied.
- Conditions 1A, 1B are not satisfied.
- Conditions 3A and 3B are partially satisfied due to the overlap in hair care preparations and the shared target market for those products.
- Given the partial overlap in goods/services and the shared target market but the lack of satisfaction in key similarity conditions (1A, 1B), the conflict grade should be moderate rather than high.

Reason:
Reasoning for Conflict:
1A - The existing trademark "MH" is not a character-for-character match with the proposed trademark "MH BY MOTHERHOOD."
1B - The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are not semantically equivalent.
1C - The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are phonetically similar due to the shared term "MH."
1D - The term "MH" is in the primary position in the proposed trademark "MH BY MOTHERHOOD." and the term "MH" is in the primary position in the existing trademark "MH." However, Existing trademark name is a single word, Hence Condition 1D is not applicable.
2  - Both the existing and proposed trademarks are in International Class 3.
3A - There is a partial overlap in the goods/services, specifically in the area of hair care preparations.
3B - There is a partial overlap in the target market, specifically for consumers interested in hair care preparations.

Conclusion:
- Conditions 1C, 2, and (1D not applicable) are satisfied.
- Conditions 1A, 1B are not satisfied.
- Conditions 3A and 3B are partially satisfied due to the overlap in hair care preparations and the shared target market for those products.
- Given the partial overlap in goods/services and the shared target market but the lack of satisfaction in key similarity conditions (1A, 1B), the conflict grade should be moderate rather than high.

- Conflict Grade: Moderate

"""         },            
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: SCOOPT'D\n
                                            Goods/Services: Hypoallergenic and vegan-friendly ice cream\n 
                                            International Class Numbers: 30\n
                                            Status: Registered\n
                                            Owner: Scoopt'd: Dairy Free Treats\n
                                            
                                            Proposed Trademark:\n
                                            Name: SCOOP-A-PALOOZA\n 
                                            Goods/Services: Ice cream\n
                                            International Class Numbers: 30\n"""
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "SCOOPT'D" is not a character-for-character match with the proposed trademark "SCOOP-A-PALOOZA."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are partially phonetically similar due to the shared term "SCOOP," which could lead to confusion. However, they are not fully phonetically similar.
- Condition 1C is not satisfied.

Step 4: Condition 1D - Primary Position
- The term "SCOOP" is in the primary position in the proposed trademark "SCOOP-A-PALOOZA."
- The term "SCOOP" is in the primary position in the existing trademark "SCOOPT'D."
- The full proposed trademark "SCOOP-A-PALOOZA" does not match the primary position of the existing trademark "SCOOPT'D."
- As if the existing trademark name is a single word, then condition 1D is not applicable.
- Existing trademark name is a single word, Hence Condition 1D is not applicable.
- Condition 1D is not applicable.

Step 5: Condition 2 - Class Overlap
- Both the existing and proposed trademarks are in International Class 30.
- Condition 2 is satisfied.

Step 6: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include hypoallergenic and vegan-friendly ice cream.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include ice cream.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- Yes, both trademarks cover ice cream.
- Condition 3A is satisfied.

Step 7: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in hypoallergenic and vegan-friendly ice cream.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in ice cream.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- Yes, both trademarks target consumers interested in ice cream.
- Condition 3B is satisfied.

Conclusion:
- Conditions 2, 3A, and 3B and (1D not applicable) are satisfied.
- Conditions 1A, 1B, 1C are not satisfied.
- Given the overlap in goods/services and the shared target market but the lack of satisfaction in key similarity conditions (1A, 1B, 1C, and 1D), the conflict grade should be moderate rather than high.

Reason:
Reasoning for Conflict:
1A - The existing trademark "SCOOPT'D" is not a character-for-character match with the proposed trademark "SCOOP-A-PALOOZA."
1B - The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are not semantically equivalent.
1C - The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are partially phonetically similar due to the shared term "SCOOP," which could lead to confusion. However, they are not fully phonetically similar.
1D - The full proposed trademark "SCOOP-A-PALOOZA" does not match the primary position of the existing trademark "SCOOPT'D." However, Existing trademark name is a single word, Hence Condition 1D is not applicable.
2  - Both the existing and proposed trademarks are in International Class 30.
3A - Both trademarks cover ice cream.
3B - Both trademarks target consumers interested in ice cream.

Conclusion:
- Conditions 2, 3A, and 3B and (1D not applicable) are satisfied.
- Conditions 1A, 1B, 1C are not satisfied.
- Given the overlap in goods/services and the shared target market but the lack of satisfaction in key similarity conditions (1A, 1B, 1C, and 1D), the conflict grade should be moderate rather than high.

- Conflict Grade: Moderate

"""
            },
            {"role": "user", "content": f"""Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: {existing_trademark['trademark_name']}\n
                                            Goods/Services: {existing_trademark['goods_services']}\n 
                                            International Class Numbers: {existing_trademark['international_class_number']}\n
                                            Status: {existing_trademark['status']}\n
                                            Owner: {existing_trademark['owner']}\n
                                            Proposed Trademark:\n
                                            Name: {proposed_name}\n 
                                            Goods/Services: {proposed_goods_services}\n
                                            International Class Numbers: {proposed_classes}\n"""
            }
        ],
        max_tokens=4000,
        temperature=0,
        top_p = 1
    )
    Treasoning = response_reasoning.choices[0].message['content'].strip()
    reasoning = Treasoning.split("Reason:", 1)[1].strip()
    conflict_grade = Treasoning.split("Conflict Grade:", 1)[1].strip() 
    progress_bar.progress(70)

    return {
        'Trademark name': existing_trademark['trademark_name'],
        'Trademark status': existing_trademark['status'],
        'Trademark owner': existing_trademark['owner'],
        'Trademark class Number': existing_trademark['international_class_number'],
        'Trademark serial number' : existing_trademark['serial_number'],
        'Trademark registration number' : existing_trademark['registration_number'],
        'Trademark design phrase' : existing_trademark['design_phrase'],
        'conflict_grade': conflict_grade,
        'reasoning': reasoning
    }


def extract_proposed_trademark_details(file_path: str) -> Dict[str, Union[str, List[int]]]:
    """ Extract proposed trademark details from the given input format """
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            
    name_match = re.search(r'Mark Searched:\s*(.*?)(?=\s*Client Name:)', page_text, re.IGNORECASE | re.DOTALL)
    if name_match:
        proposed_details["proposed_trademark_name"] = name_match.group(1).strip()

    goods_services_match = re.search(r'Goods/Services:\s*(.*?)(?=\s*Trademark Research Report)', page_text, re.IGNORECASE | re.DOTALL)
    if goods_services_match:
        proposed_details["proposed_goods_services"] = goods_services_match.group(1).strip()
    
    # Use LLM to find the international class number based on goods & services
    if "proposed_goods_services" in proposed_details:
        goods_services = proposed_details["proposed_goods_services"]
        class_numbers = find_class_numbers(goods_services)
        proposed_details["proposed_nice_classes_number"] = class_numbers
    
    return proposed_details

def find_class_numbers(goods_services: str) -> List[int]:
    """ Use LLM to find the international class numbers based on goods & services """
        # Initialize AzureChatOpenAI
    azure_endpoint = "https://chat-gpt-a1.openai.azure.com/"
    api_key = "c09f91126e51468d88f57cb83a63ee36"
    deployment_name = "DanielChatGPT16k"

    openai.api_type = "azure"
    openai.api_key = api_key
    openai.api_base = azure_endpoint
    openai.api_version = "2024-02-15-preview"
        
    response = openai.ChatCompletion.create(
        engine=deployment_name,
        messages=[
            {"role": "system", "content": "You are a helpful assistant for finding the International class number of provided Goods & Services."},
            {"role": "user", "content": "The goods/services are: IC 003: SKIN CARE PREPARATIONS; COSMETICS; BABY CARE PRODUCTS, NAMELY, SKIN SOAPS, BABY WASH, BABY BUBBLE BATH, BABY LOTIONS, BABY SHAMPOOS; SKIN CLEANSERS; BABY WIPES; NON− MEDICATED DIAPER RASH OINTMENTS AND LOTIONS; SKIN LOTIONS, CREAMS, MOISTURIZERS, AND OILS; BODY WASH; BODY SOAP; DEODORANTS; PERFUME; HAIR CARE PREPARATIONS. Find the international class numbers."},
            {"role": "assistant", "content": "The international class numbers : 03"},
            {"role": "user", "content": "The goods/services are: LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES. Find the international class numbers."},
            {"role": "assistant", "content": "The international class numbers : 18,35"},
            {"role": "user", "content": f"The goods/services are: {goods_services}. Find the international class numbers."}
        ],
        max_tokens=150,
        temperature=0
    )
    class_numbers_str = response.choices[0].message['content'].strip()
    
    # Extracting class numbers and removing duplicates
    class_numbers = re.findall(r'(?<!\d)\d{2}(?!\d)', class_numbers_str)  # Look for two-digit numbers
    class_numbers = ','.join(set(class_numbers))  # Convert to set to remove duplicates, then join into a single string
    
    return class_numbers

def extract_proposed_trademark_details2(file_path: str) -> Dict[str, Union[str, List[int]]]:
    """ Extract proposed trademark details from the first page of the document """
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            
            name_match = re.search(r'Name:\s*(.*?)(?=\s*Nice Classes:)', page_text)
            if name_match:
                proposed_details["proposed_trademark_name"] = name_match.group(1).strip()
                
            nice_classes_match = re.search(r'Nice Classes:\s*(\d+(?:,\s*\d+)*)', page_text)
            if nice_classes_match:
                proposed_details["proposed_nice_classes_number"] = nice_classes_match.group(1).strip()
            
            goods_services_match = re.search(r'Goods & Services:\s*(.*?)(?=\s*Registers|$)', page_text, re.IGNORECASE | re.DOTALL)
            if goods_services_match:
                proposed_details["proposed_goods_services"] = goods_services_match.group(1).strip()
    
    return proposed_details

def list_conversion(proposed_class: str) -> List[int]:
    
    azure_endpoint = "https://chat-gpt-a1.openai.azure.com/"
    api_key = "c09f91126e51468d88f57cb83a63ee36"
    deployment_name = "DanielChatGPT16k"

    openai.api_type = "azure"
    openai.api_key = api_key
    openai.api_base = azure_endpoint
    openai.api_version = "2024-02-15-preview"

    response = openai.ChatCompletion.create(
                engine=deployment_name,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant for converting the class number string into python list of numbers.\n Respond only with python list. Example : [18,35]"},
                    {"role": "user", "content": "The class number are: 15,89. convert the string into python list of numbers."},
                    {"role": "assistant", "content": "[15,89]"},
                    {"role": "user", "content": f"The class number are: {proposed_class}. convert the string into python list of numbers."}
                ],
                max_tokens=150,
                temperature=0
            )
    lst_class = response.choices[0].message['content'].strip()
    class_value = ast.literal_eval(lst_class)
            
    return class_value

# Streamlit App  
st.title("Trademark Document Parser")  
  
# File upload  
uploaded_files = st.sidebar.file_uploader("Choose PDF files", type="pdf", accept_multiple_files=True)  
  
if uploaded_files:  
    if st.sidebar.button("Check Conflicts", key="check_conflicts"):  
        total_files = len(uploaded_files)  
        progress_bar = st.progress(0)  
        for i, uploaded_file in enumerate(uploaded_files):  
            # Save uploaded file to a temporary file path  
            temp_file_path = f"temp_{uploaded_file.name}"  
            with open(temp_file_path, "wb") as f:  
                f.write(uploaded_file.read())  
            
            sp = True
            proposed_trademark_details = extract_proposed_trademark_details(temp_file_path)  
                            
            if proposed_trademark_details:  
                proposed_name = proposed_trademark_details.get('proposed_trademark_name', 'N')  
                proposed_class = proposed_trademark_details.get('proposed_nice_classes_number')  
                proposed_goods_services = proposed_trademark_details.get('proposed_goods_services', 'N') 
                if (proposed_goods_services != 'N'): 
                    with st.expander(f"Proposed Trademark Details for {uploaded_file.name}"):  
                            st.write(f"Proposed Trademark name: {proposed_name}")  
                            st.write(f"Proposed class-number: {proposed_class}")  
                            st.write(f"Proposed Goods & Services: {proposed_goods_services}") 
                    class_list = list_conversion(proposed_class) 
                else :
                    st.write("______________________________________________________________________________________________________________________________")
                    st.write(f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}")
                    st.write("______________________________________________________________________________________________________________________________")
                    sp = False
            else:  
                
                proposed_trademark_details = extract_proposed_trademark_details2(temp_file_path)  
                
                if proposed_trademark_details:  
                    proposed_name = proposed_trademark_details.get('proposed_trademark_name', 'N')  
                    proposed_class = proposed_trademark_details.get('proposed_nice_classes_number')  
                    proposed_goods_services = proposed_trademark_details.get('proposed_goods_services', 'N')  
                    if (proposed_goods_services != 'N'): 
                        with st.expander(f"Proposed Trademark Details for {uploaded_file.name}"):  
                                st.write(f"Proposed Trademark name: {proposed_name}")  
                                st.write(f"Proposed class-number: {proposed_class}")  
                                st.write(f"Proposed Goods & Services: {proposed_goods_services}") 
                        class_list = list_conversion(proposed_class)  
                    else :
                        st.write("______________________________________________________________________________________________________________________________")
                        st.write(f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}")
                        st.write("______________________________________________________________________________________________________________________________")
                        sp = False
                else :  
                    st.error(f"Unable to extract Proposed Trademark Details for {uploaded_file.name}") 
                    sp = False 
                    continue  
            
            if (sp):    
                for i in range(1,21):
                    time.sleep(0.5)
                    progress_bar.progress(i)
                    
                progress_bar.progress(25)
                # Initialize AzureChatOpenAI
                azure_endpoint = "https://chat-gpt-a1.openai.azure.com/"
                api_key = "c09f91126e51468d88f57cb83a63ee36"
                deployment_name = "DanielChatGPT16k"

                openai.api_type = "azure"
                openai.api_key = api_key
                openai.api_base = azure_endpoint
                openai.api_version = "2024-02-15-preview"
                
                existing_trademarks = parse_trademark_details(temp_file_path)
                for i in range(25,46):
                    time.sleep(0.5)
                    progress_bar.progress(i)  
                    
                progress_bar.progress(50)
                st.success(f"Existing Trademarks Data Extracted Successfully for {uploaded_file.name}!")  
                # Display extracted details              
                
                nfiltered_list = []
                
                # Iterate over each JSON element in trademark_name_list  
                for json_element in existing_trademarks:  
                    class_numbers = json_element["international_class_number"]  
                # Check if any of the class numbers are in class_list  
                    if any(number in class_list for number in class_numbers):  
                        nfiltered_list.append(json_element)
                    
                existing_trademarks = nfiltered_list
                     
                azure_endpoint = "https://danielingitaraj.openai.azure.com/"
                api_key = "a5c4e09a50dd4e13a69e7ef19d07b48c"
                deployment_name = "GPT4"  

                openai.api_type = "azure"
                openai.api_key = api_key
                openai.api_base = azure_endpoint 
                openai.api_version = "2024-02-01"

                high_conflicts = []
                moderate_conflicts = []
                low_conflicts = []
                
                lt = len(existing_trademarks)
                
                for existing_trademark in existing_trademarks:  
                    conflict = compare_trademarks(existing_trademark, proposed_name, proposed_class, proposed_goods_services)  
                    if conflict['conflict_grade'] == "High":  
                        high_conflicts.append(conflict)  
                    elif conflict['conflict_grade'] == "Moderate":  
                        moderate_conflicts.append(conflict)  
                    else:  
                        low_conflicts.append(conflict)  
    
                st.sidebar.write("_________________________________________________")
                st.sidebar.subheader("\n\nConflict Grades : \n")  
                st.sidebar.markdown(f"File: {proposed_name}")  
                st.sidebar.markdown(f"Total number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(low_conflicts)}")
                st.sidebar.markdown(f"High Conflicts: {len(high_conflicts)}")  
                st.sidebar.markdown(f"Moderate Conflicts: {len(moderate_conflicts)}")  
                st.sidebar.markdown(f"Low Conflicts: {len(low_conflicts)}")  
                st.sidebar.write("_________________________________________________")
    
                document = Document()  
                
                document.add_heading(f'Trademark Conflict List for {proposed_name} :')            
                document.add_paragraph(f"\n\nTotal number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(low_conflicts)}\n- High Conflicts: {len(high_conflicts)}\n- Moderate Conflicts: {len(moderate_conflicts)}\n- Low Conflicts: {len(low_conflicts)}\n")  
                
                if len(high_conflicts) > 0:  
                            document.add_heading('Trademarks with High Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_high = pd.DataFrame(high_conflicts) 
                            df_high = df_high.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])  
                            # Create a table in the Word document    
                            table_high = document.add_table(df_high.shape[0] + 1, df_high.shape[1])
                            # Set a predefined table style (with borders)  
                            table_high.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_high.columns):  
                                table_high.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_high.iterrows():  
                                for j, value in enumerate(row):  
                                    table_high.cell(i + 1, j).text = str(value)

                if len(moderate_conflicts) > 0:  
                            document.add_heading('Trademarks with Moderate Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_moderate = pd.DataFrame(moderate_conflicts)
                            df_moderate = df_moderate.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])  
                            # Create a table in the Word document    
                            table_moderate = document.add_table(df_moderate.shape[0] + 1, df_moderate.shape[1])
                            # Set a predefined table style (with borders)  
                            table_moderate.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_moderate.columns):  
                                table_moderate.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_moderate.iterrows():  
                                for j, value in enumerate(row):  
                                    table_moderate.cell(i + 1, j).text = str(value)

                if len(low_conflicts) > 0:  
                            document.add_heading('Trademarks with Low Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_low = pd.DataFrame(low_conflicts)  
                            df_low = df_low.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])
                            # Create a table in the Word document    
                            table_low = document.add_table(df_low.shape[0] + 1, df_low.shape[1])
                            # Set a predefined table style (with borders)  
                            table_low.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_low.columns):  
                                table_low.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_low.iterrows():  
                                for j, value in enumerate(row):  
                                    table_low.cell(i + 1, j).text = str(value)
                            
                def add_conflict_paragraph(document, conflict):  
                    p = document.add_paragraph(f"Trademark Name : {conflict.get('Trademark name', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Status : {conflict.get('Trademark status', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Owner : {conflict.get('Trademark owner', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)  
                    p = document.add_paragraph(f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark registration number : {conflict.get('Trademark registration number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0) 
                    p = document.add_paragraph(f"Trademark Design phrase : {conflict.get('Trademark design phrase', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0) 
                    p = document.add_paragraph(f"{conflict.get('reasoning','N/A')}\n")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                
                if len(high_conflicts) > 0:  
                    document.add_heading('Trademarks with High Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in high_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                
                if len(moderate_conflicts) > 0:  
                    document.add_heading('Trademarks with Moderate Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in moderate_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                
                if len(low_conflicts) > 0:  
                    document.add_heading('Trademarks with Low Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in low_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                        
                for i in range(70,96):
                    time.sleep(0.5)
                    progress_bar.progress(i)  
                    
                progress_bar.progress(100)
    
                filename = proposed_name
                doc_stream = BytesIO()  
                document.save(doc_stream)  
                doc_stream.seek(0)  
                download_table = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_stream.read()).decode()}" download="{filename + " Trademark Conflict Report"}.docx">Download: {filename}</a>'  
                st.sidebar.markdown(download_table, unsafe_allow_html=True)  
                st.success(f"{proposed_name} Document conflict report successfully completed!")
                st.write("______________________________________________________________________________________________________________________________")
  
        progress_bar.progress(100)
        st.success("All documents processed successfully!")  
